import docx, os
os.environ['PYTHONIOENCODING'] = 'utf-8'

doc = docx.Document(r'C:\Users\malware\Desktop\final_thesis\report_thesis\MEng_Thesis_Final.docx')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

for i, sec in enumerate(doc.sections):
    print(f'=== Section {i} ===')
    # Check different footer types
    footer = sec.footer
    print(f'  Default footer paragraphs: {len(footer.paragraphs)}')
    for j, p in enumerate(footer.paragraphs):
        txt = p.text[:80] if p.text else '(empty)'
        fld = p._element.findall('.//' + W + 'fldChar')
        instr = [it.text or '' for it in p._element.findall('.//' + W + 'instrText')]
        print(f'    [{j}] "{txt}"  fields={len(fld)}  instr={instr}')
    
    # Try to get first page footer
    try:
        first_footer = sec.first_page_footer
        if first_footer is not None:
            print(f'  First page footer: {len(first_footer.paragraphs)} paragraphs')
            for j, p in enumerate(first_footer.paragraphs):
                txt = p.text[:80] if p.text else '(empty)'
                print(f'    [{j}] "{txt}"')
    except:
        print('  No first page footer')

    # Check even page footer
    try:
        even_footer = sec.even_page_footer
        if even_footer is not None:
            print(f'  Even page footer: {len(even_footer.paragraphs)} paragraphs')
    except:
        print('  No even page footer')
    
    # Check relationships
    sectPr = sec._sectPr
    fr = sectPr.findall(W + 'footerReference')
    for fref in fr:
        rId = fref.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        ftype = fref.get(W + 'type')
        print(f'  footerReference: type={ftype} rId={rId}')
