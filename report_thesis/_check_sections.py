import docx
from docx.oxml.ns import qn

doc = docx.Document(r'C:\Users\malware\Desktop\final_thesis\report_thesis\MEng_Thesis_Final.docx')
print(f'Total sections: {len(doc.sections)}')
for i, sec in enumerate(doc.sections):
    print(f'Section {i}:')
    print(f'  start_type={sec.start_type}')
    # Check pgNumType in sectPr
    sectPr = sec._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        print(f'  pgNumType fmt={pgNumType.get(qn("w:fmt"))} start={pgNumType.get(qn("w:start"))}')
    else:
        print(f'  pgNumType: none')
    # Check footer reference
    footerRef = sectPr.find(qn('w:footerReference'))
    if footerRef is not None:
        print(f'  footerRef type={footerRef.get(qn("w:type"))}')
    print(f'  Footer paragraphs: {len(sec.footer.paragraphs)}')
    for j, p in enumerate(sec.footer.paragraphs):
        txt = p.text[:80] if p.text else '(empty)'
        print(f'    Footer[{j}]: "{txt}"')
    # Check for PAGE field codes in footer
    ns = qn('w:')
    for j, p in enumerate(sec.footer.paragraphs):
        fld_chars = p._element.findall('.//' + ns + 'fldChar')
        instr = p._element.findall('.//' + ns + 'instrText')
        if fld_chars:
            print(f'    Footer[{j}] HAS FIELD CODES')
        for instr_text in instr:
            if instr_text.text:
                print(f'    Footer[{j}] instrText: {instr_text.text[:60]}')
