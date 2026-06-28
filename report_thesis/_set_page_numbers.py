"""
Set up page numbering in MEng_Thesis_Final.docx:
  - Cover page (title page): NO page number (separate section, no footer)
  - Front matter (Declaration through List of Abbreviations): Roman numerals (i, ii, iii...), centered
  - Chapter 1 onwards: Arabic numerals (1, 2, 3...), centered, restarting at 1
"""
import docx
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import copy
import os

os.environ['PYTHONIOENCODING'] = 'utf-8'

DOCX_PATH = r'C:\Users\malware\Desktop\final_thesis\report_thesis\MEng_Thesis_Final.docx'

doc = docx.Document(DOCX_PATH)
body = doc.element.body
W = qn('w:')

# --- Helper: create a PAGE field paragraph XML ---
def make_page_field_para():
    """Create a paragraph containing just the PAGE field, to be centered."""
    xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:jc w:val="center"/>'
        '<w:r>'
        '  <w:fldChar w:fldCharType="begin"/>'
        '</w:r>'
        '<w:r>'
        '  <w:instrText> PAGE </w:instrText>'
        '</w:r>'
        '<w:r>'
        '  <w:fldChar w:fldCharType="separate"/>'
        '</w:r>'
        '<w:r>'
        '  <w:t>1</w:t>'
        '</w:r>'
        '<w:r>'
        '  <w:fldChar w:fldCharType="end"/>'
        '</w:r>'
        '</w:p>'
    )
    return etree.fromstring(xml)

def make_empty_footer_para():
    """Create an empty footer paragraph."""
    xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:t></w:t></w:r>'
        '</w:p>'
    )
    return etree.fromstring(xml)

# --- Step 1: Find key paragraph indices ---
paras = list(body.iterchildren(W + 'p'))

# Find the title page last para (copyright at P52)
copyright_idx = None
for i, p in enumerate(paras):
    texts = p.findall('.//' + W + 't')
    tc = ''.join(t.text or '' for t in texts).strip()
    if 'All rights reserved' in tc:
        copyright_idx = i
        break
print(f'Copyright paragraph at index: {copyright_idx}')

# Find the CHAPTER 1 paragraph (where body section starts)
ch1_idx = None
for i, p in enumerate(paras):
    texts = p.findall('.//' + W + 't')
    tc = ''.join(t.text or '' for t in texts).strip()
    if tc == 'CHAPTER 1: INTRODUCTION':
        ch1_idx = i
        break
print(f'CHAPTER 1 at index: {ch1_idx}')

# Find the existing inline sectPr (currently at the paragraph before CHAPTER 1)
existing_sectPr_para_idx = None
for i, p in enumerate(paras):
    pPr = p.find(W + 'pPr')
    if pPr is not None:
        sp = pPr.find(W + 'sectPr')
        if sp is not None:
            existing_sectPr_para_idx = i
            break
print(f'Existing inline sectPr at paragraph index: {existing_sectPr_para_idx}')

# Also find the final sectPr
final_sectPr = None
for child in body.iterchildren():
    if child.tag == W + 'sectPr':
        final_sectPr = child
        break  # the last one
print(f'Final sectPr found: {final_sectPr is not None}')

# --- Step 2: Verify which section is which ---
print()
print('=== Current Section Structure ===')
for i, sec in enumerate(doc.sections):
    sp = sec._sectPr
    pgNumType = sp.find(W + 'pgNumType')
    fmt = pgNumType.get(W + 'fmt') if pgNumType is not None else None
    start = pgNumType.get(W + 'start') if pgNumType is not None else None
    print(f'Section {i}: fmt={fmt}, start={start}')
    # Check footer
    footer = sec.footer
    for j, p in enumerate(footer.paragraphs):
        txt = p.text[:60] if p.text else '(empty)'
        print(f'  Footer[{j}]: "{txt}"')

# === STEP 3: Insert section break after title page ===
# Add an inline sectPr to the copyright paragraph (paragraph index copyright_idx)
# This creates a new section starting at DECLARATION
copyright_para = paras[copyright_idx]
pPr = copyright_para.find(W + 'pPr')
if pPr is None:
    pPr = etree.SubElement(copyright_para, W + 'pPr')
    # Insert at beginning
    copyright_para.insert(0, pPr)

# Check if there's already a sectPr in this paragraph
existing_sp = pPr.find(W + 'sectPr')
if existing_sp is not None:
    pPr.remove(existing_sp)
    print('Removed existing sectPr from copyright paragraph')

# Create new sectPr for the title page section (no footer)
new_sectPr = etree.SubElement(pPr, W + 'sectPr')
new_sectPr.set(W + 'type', 'nextPage')

# Add pgNumType with no formatting (no page numbers)
pgNumType = etree.SubElement(new_sectPr, W + 'pgNumType')
pgNumType.set(W + 'fmt', 'none')

# Add empty footer reference (so title page has no footer)
# We don't add a footerReference at all, which means no footer

# Add page size and margins (copy from existing section properties to maintain layout)
# Actually, copy from the existing section's sectPr
existing_first_sp = paras[existing_sectPr_para_idx].find(W + 'pPr/' + W + 'sectPr')
if existing_first_sp is not None:
    for child in existing_first_sp.iterchildren():
        if child.tag in [W + 'pgSz', W + 'pgMar']:
            new_sectPr.append(copy.deepcopy(child))

print(f'Added section break at paragraph {copyright_idx} (title page end)')

# === STEP 4: Modify the existing section break at P312 (front matter → body) ===
# This paragraph already has an inline sectPr. We need to:
# 1. Set pgNumType to lowerRoman for the front matter section
# 2. Set up the footer for the front matter section

# Find the existing sectPr at para existing_sectPr_para_idx
existing_p = paras[existing_sectPr_para_idx]
existing_pPr = existing_p.find(W + 'pPr')
existing_sp_inline = existing_pPr.find(W + 'sectPr') if existing_pPr is not None else None

if existing_sp_inline is not None:
    # Remove existing pgNumType and add lowerRoman
    old_pgNum = existing_sp_inline.find(W + 'pgNumType')
    if old_pgNum is not None:
        existing_sp_inline.remove(old_pgNum)
    new_pgNum = etree.SubElement(existing_sp_inline, W + 'pgNumType')
    new_pgNum.set(W + 'fmt', 'lowerRoman')
    print(f'Set sectPr at para {existing_sectPr_para_idx} to lowerRoman')

# === STEP 5: Modify final sectPr (body section) ===
if final_sectPr is not None:
    # Remove existing pgNumType
    old_pgNum = final_sectPr.find(W + 'pgNumType')
    if old_pgNum is not None:
        final_sectPr.remove(old_pgNum)
    # Add new pgNumType with decimal starting at 1
    new_pgNum = etree.SubElement(final_sectPr, W + 'pgNumType')
    new_pgNum.set(W + 'fmt', 'decimal')
    new_pgNum.set(W + 'start', '1')
    print('Set final sectPr to decimal starting at 1')

# === STEP 6: Update footers for each section ===
# Re-read sections after structural changes
# The document now has 3 sections: title, front matter, body

# --- Helper to set footer content ---
def set_footer_content(section, page_field=True):
    """Clear footer and set either a PAGE field or empty content."""
    footer = section.footer
    # Remove existing paragraphs
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    
    if page_field:
        new_para = make_page_field_para()
        footer._element.append(new_para)
    else:
        new_para = make_empty_footer_para()
        footer._element.append(new_para)

# Process each section
for i, sec in enumerate(doc.sections):
    sp = sec._sectPr
    pgNumType = sp.find(W + 'pgNumType')
    fmt = pgNumType.get(W + 'fmt') if pgNumType is not None else None
    start = pgNumType.get(W + 'start') if pgNumType is not None else None
    
    if i == 0:
        # Title page section: no page number
        set_footer_content(sec, page_field=False)
        # Also ensure different first page isn't interfering
        sp.set(W + 'titlePg', '1')
        print(f'Section {i} (title): EMPTY footer')
    elif i == 1:
        # Front matter section: Roman numerals
        set_footer_content(sec, page_field=True)
        print(f'Section {i} (front matter): Roman numeral footer')
    else:
        # Body section: Arabic numerals starting at 1
        set_footer_content(sec, page_field=True)
        print(f'Section {i} (body): Arabic numeral footer starting at 1')

# === STEP 7: Link the title section footer as NOT linked to previous ===
# By default, sections inherit from previous. We need to unlink for title page.
# In OOXML, this means the footer reference should NOT have a 'link' attribute,
# and we should set w:titlePg on the section.

# Actually, we already handled this. Let's verify.

# === STEP 8: Save ===
doc.save(DOCX_PATH)
print('\nSaved!')

# === VERIFY ===
print('\n=== New Section Structure ===')
doc2 = docx.Document(DOCX_PATH)
for i, sec in enumerate(doc2.sections):
    sp = sec._sectPr
    pgNumType = sp.find(W + 'pgNumType')
    fmt = pgNumType.get(W + 'fmt') if pgNumType is not None else None
    start = pgNumType.get(W + 'start') if pgNumType is not None else None
    print(f'Section {i}: fmt={fmt}, start={start}')
    footer = sec.footer
    for j, p in enumerate(footer.paragraphs):
        txt = p.text[:60] if p.text else '(empty)'
        # Check for field codes
        fld = p._element.findall('.//' + W + 'fldChar')
        instr = p._element.findall('.//' + W + 'instrText')
        has_page = any((it.text or '').strip() == 'PAGE' for it in instr) if instr else False
        print(f'  Footer[{j}]: "{txt}"  fields={len(fld)}  has_PAGE={has_page}')
