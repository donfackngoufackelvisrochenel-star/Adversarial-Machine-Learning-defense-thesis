import docx
import sys

doc = docx.Document('report_thesis/MEng_Thesis_Final.docx')

# Extract tables
for i, table in enumerate(doc.tables):
    print(f"\n{'='*60}")
    print(f"TABLE {i} ({len(table.rows)} rows x {len(table.columns)} cols)")
    print(f"{'='*60}")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))

# Also extract the result paragraphs around chapter 4
print(f"\n\n{'='*60}")
print("RESULT PARAGRAPHS")
print(f"{'='*60}")
in_results = False
for p in doc.paragraphs:
    text = p.text.strip()
    if '4.1.1 Clean Performance' in text:
        in_results = True
    if '4.2 Discussion' in text:
        in_results = False
    if in_results and text:
        try:
            print(text[:500])
        except:
            print(text.encode('ascii', errors='replace').decode()[:500])
