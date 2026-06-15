from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

os.chdir(r'C:\Users\malware\Desktop\final_thesis')

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)

# ============================================================
doc.add_heading('AML DEFENSE PROJECT — COMPLETE DOCUMENTATION', level=0)

doc.add_paragraph(
    'This document explains every component of the project. Read it section by section '
    'to understand how the code is organised and how all the pieces fit together.'
)

# ============================================================
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'This project implements a complete pipeline for training, attacking, and defending '
    'machine learning classifiers on the CIC_IoMT_2024 dataset (Internet of Medical Things '
    'network traffic data, 51 classes). The goal is to evaluate how vulnerable ML models are '
    'to evasion attacks (FGSM and PGD) and to compare four defense strategies.'
)
doc.add_paragraph('Key components:', style='List Bullet')
for b in [
    'Three ML classifiers: Random Forest, XGBoost, LightGBM',
    'Two evasion attacks: FGSM (single-step), PGD (iterative)',
    'Four defense strategies: Adversarial Training, Feature Selection, Ensemble, Feature Squeezing',
    'Interactive dashboard (Streamlit) for visual exploration',
    'REST API (FastAPI) for remote inference',
    'SHAP explanations for model interpretability',
]:
    doc.add_paragraph(b, style='List Bullet')

# ============================================================
doc.add_heading('2. Project Structure', level=1)
code = (
    'final_thesis/\n'
    '├── configs/config.py              # Central configuration\n'
    '├── data/\n'
    '│   ├── raw/                       # Place your dataset here\n'
    '│   └── processed/                 # Auto-generated\n'
    '├── src/\n'
    '│   ├── preprocessing/loader.py    # Load, clean, split data\n'
    '│   ├── preprocessing/processor.py # Z-score scaling & label encoding\n'
    '│   ├── features/builder.py        # Feature importance selection\n'
    '│   ├── models/trainer.py          # RF, XGBoost, LightGBM training\n'
    '│   ├── attacks/fgsm.py            # Fast Gradient Sign Method\n'
    '│   ├── attacks/pgd.py             # Projected Gradient Descent\n'
    '│   ├── defenses/adversarial_training.py\n'
    '│   ├── defenses/feature_selection.py\n'
    '│   ├── defenses/ensemble.py\n'
    '│   ├── defenses/feature_squeezing.py\n'
    '│   ├── evaluation/metrics.py      # Accuracy, Precision, Recall, F1, ASR\n'
    '│   ├── evaluation/comparison.py   # Robustness curves & reports\n'
    '│   ├── explainability/explainer.py# SHAP explanations\n'
    '│   ├── utils/helpers.py           # Logging, save/load\n'
    '│   ├── api/main.py                # FastAPI endpoints\n'
    '│   └── dashboard/streamlit_app.py # Interactive dashboard\n'
    '├── models/                        # Saved model pickles\n'
    '├── reports/                       # Plots, CSVs, metrics JSON\n'
    '├── logs/                          # Pipeline logs\n'
    '├── scripts/run_pipeline.py        # End-to-end pipeline\n'
    '├── deployment/\n'
    '│   ├── Dockerfile\n'
    '│   └── docker-compose.yml\n'
    '├── tests/                         # Unit tests\n'
    '└── requirements.txt\n'
)
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ============================================================
doc.add_heading('3. Configuration (configs/config.py)', level=1)
doc.add_paragraph(
    'All tunable parameters are defined in one central file.'
)
doc.add_heading('Data Splitting', level=2)
doc.add_paragraph(
    'TARGET_COLUMN = "label"\n'
    'TEST_SIZE = 0.2  (20% held out for final testing)\n'
    'VAL_SIZE = 0.1   (10% of the training portion = 8% of total for validation)\n'
    'RANDOM_STATE = 42\n\n'
    'Final split: 72% train, 8% validation, 20% test.\n\n'
    'The validation set is carved FROM the training set (10% of 80% = 8% of total). '
    'This gives you 80% for training-related work (72% actual train + 8% validation for tuning) '
    'and 20% held out as a completely unseen test set for final evaluation.'
)
doc.add_heading('Other Parameters', level=2)
doc.add_paragraph(
    'MAX_ROWS = 300000 — caps dataset size for memory safety.\n'
    'CHUNKSIZE = 50000 — rows per chunk for CSV loading.\n\n'
    'Model hyperparameters: RF (30 trees, depth 8), XGBoost (30 trees, depth 4, lr=0.1), '
    'LightGBM (same as XGBoost).\n\n'
    'Attack: FGSM (epsilon=0.1), PGD (epsilon=0.1, alpha=0.01, iter=10).\n\n'
    'Defense: adv_training (epsilon=0.1, 1 epoch), feature_selection (top_k=20), '
    'ensemble (soft voting), feature_squeezing (bit_depth=4).'
)

# ============================================================
doc.add_heading('4. Data Pipeline (src/preprocessing/)', level=1)
doc.add_heading('4.1 loader.py', level=2)
doc.add_paragraph(
    'find_dataset() — Scans data/raw/ for CSV, TXT, ZIP, GZ, or Parquet files. '
    'Extracts archives automatically. Prefers Parquet (fastest).\n\n'
    'load_data() — Main loading function. Supports chunked reading. When max_rows is set, '
    'loads a larger pool (max_rows x 10) for class diversity, then randomly samples down.\n\n'
    'clean_data() — Removes duplicates, drops all-NaN columns, forward-fills missing values, '
    'converts float16 to float32.\n\n'
    'split_data() — Finds the target column ("label", "class", "target", or last column). '
    'Drops features with >0.90 correlation with label (prevents trivial perfect classification). '
    'Encodes string labels to integers. One-hot encodes categorical features. '
    'Splits into 72% train, 8% validation, 20% test using stratified sampling.'
)
doc.add_heading('4.2 processor.py', level=2)
doc.add_paragraph(
    'DataProcessor wraps StandardScaler (z-score: x\' = (x - mean) / std) and LabelEncoder. '
    'Usage: fit on training data, then transform validation and test sets with the same parameters. '
    'Supports save/load via joblib.'
)

# ============================================================
doc.add_heading('5. Feature Engineering (src/features/builder.py)', level=1)
doc.add_paragraph(
    'select_features_by_importance() — Trains a Random Forest probe (50 trees), '
    'returns the top-K features by importance.\n\n'
    'reduce_to_features() — Subsets a DataFrame to specified columns.'
)

# ============================================================
doc.add_heading('6. Models (src/models/trainer.py)', level=1)
doc.add_paragraph(
    'MODEL_REGISTRY maps names to classes: random_forest -> RandomForestClassifier, '
    'xgboost -> XGBClassifier, lightgbm -> LGBMClassifier.\n\n'
    'train_model(name, X_train, y_train, X_val, y_val) — Instantiates model with config '
    'params, fits, prints train/val metrics.\n\n'
    'train_all_models() — Trains all three, saves each as {name}_clean.pkl, returns dict.\n\n'
    'save_model() / load_model() — joblib serialisation.'
)

# ============================================================
doc.add_heading('7. Attacks (src/attacks/)', level=1)
doc.add_heading('7.1 FGSM', level=2)
doc.add_paragraph(
    'Single-step attack. Perturbs each sample in the direction that maximally increases loss.\n\n'
    'Adaptation for tree models: Uses feature importances + prediction confidence as gradient proxy. '
    'Perturbation scaled by feature standard deviation.\n\n'
    'Algorithm: For each sample, if correctly classified, push in the direction that maximises loss; '
    'if already misclassified, apply random perturbation. Clips to feature bounds if provided.'
)
doc.add_heading('7.2 PGD', level=2)
doc.add_paragraph(
    'Iterative attack (default 10 steps). Starts from random noise inside epsilon-ball, '
    'takes small steps (alpha=0.01), and projects back after each iteration. '
    'Produces stronger adversarial examples than FGSM.'
)

# ============================================================
doc.add_heading('8. Defenses (src/defenses/)', level=1)
doc.add_heading('8.1 Adversarial Training', level=2)
doc.add_paragraph(
    'Augments training set with adversarial examples. For each epoch, generates FGSM examples '
    'from current model, appends to training pool, retrains from scratch. '
    'Trains model to be robust against the specific attack used during augmentation.'
)
doc.add_heading('8.2 Feature Selection', level=2)
doc.add_paragraph(
    'Ranks features by vulnerability to noise (accuracy drop when feature is perturbed). '
    'Keeps only the K least vulnerable features (default K=20). '
    'Reduces attack surface by removing features adversaries can exploit.'
)
doc.add_heading('8.3 Ensemble', level=2)
doc.add_paragraph(
    'Combines RF + XGBoost + LightGBM via soft voting (averages class probabilities). '
    'An adversary must fool all three models simultaneously to change the output.'
)
doc.add_heading('8.4 Feature Squeezing', level=2)
doc.add_paragraph(
    'Bit-depth reduction: quantises features to 2^bits levels (default 4 bits = 16 levels). '
    'Median smoothing: replaces each value with median of neighbouring rows. '
    'Both remove small perturbations while preserving legitimate classification.'
)

# ============================================================
doc.add_heading('9. Evaluation (src/evaluation/)', level=1)
doc.add_heading('9.1 metrics.py', level=2)
doc.add_paragraph(
    'compute_metrics() returns: accuracy, precision, recall, f1_score, (roc_auc for binary).\n\n'
    'attack_success_rate() = fraction of adversarial examples causing misclassification.\n'
    'robustness_curve() = accuracy at each epsilon.\n'
    'summarize_defense_comparison() = DataFrame of defense vs metrics.'
)
doc.add_heading('9.2 comparison.py', level=2)
doc.add_paragraph(
    'compare_attack_defense() — Evaluates all clean + defended models across epsilon values.\n'
    'plot_robustness_comparison() — Saves accuracy-vs-epsilon plot.\n'
    'generate_report_table() — Pivot table (defenses x epsilons) saved to CSV.'
)

# ============================================================
doc.add_heading('10. Explainability (src/explainability/explainer.py)', level=1)
doc.add_paragraph(
    'Uses SHAP TreeExplainer (falls back to KernelExplainer). '
    'Generates summary bar plots showing feature importance. '
    'Handles multi-class by averaging absolute SHAP values across classes.'
)

# ============================================================
doc.add_heading('11. Pipeline (scripts/run_pipeline.py)', level=1)
steps = [
    'Step 1: LOAD & PROCESS — load dataset, clean, split 72/8/20, scale/encode',
    'Step 2: TRAIN — train RF, XGBoost, LightGBM, save to models/',
    'Step 3: EVALUATE — compute metrics, confusion matrices, save to reports/',
    'Step 4: ATTACK — run FGSM + PGD at multiple epsilon values, report ASR',
    'Step 5: DEFEND — adversarial training, feature selection, ensemble, squeezing',
    'Step 6: REPORT — robustness curves, comparison tables',
    'Step 7: EXPLAIN — SHAP values, save summary plots',
]
for s in steps:
    doc.add_paragraph(s)

doc.add_paragraph(
    'Flags: --skip-attack, --skip-defense, --skip-explain, --max-rows N, --chunksize N\n'
    'Run: python scripts/run_pipeline.py'
)

# ============================================================
doc.add_heading('12. Dashboard (src/dashboard/streamlit_app.py)', level=1)
doc.add_paragraph(
    'Streamlit GUI with authentication. Sidebar: logout button, dataset selection, loading, '
    'training, attacks, defenses. '
    'Main panel: dataset overview, pipeline audit, model performance table, best algorithm '
    'banner (with Accuracy, Precision, Recall, F1), metric comparison chart, per-class '
    'confusion matrix, attack results, defense comparison, SHAP feature importance.\n\n'
    'Auto-Loading: When pre-trained models and processor.pkl exist on disk, the dashboard '
    'automatically loads them. The processor is applied to scale raw CSV data on load, '
    'ensuring metrics computed on the test set are consistent with how the models were trained.\n\n'
    'Authentication: login page (username/password) gates all content. Credentials in '
    'configs/config.py (default: admin / ciciomt2024). Logout in sidebar.\n\n'
    'Live Inference Demo section at the bottom:\n'
    '- Model selector, FGSM attack toggle, epsilon control, delay slider\n'
    '- "Next Sample" button streams one test sample and runs prediction\n'
    '- Running accuracy tracker, average latency display\n'
    '- Accuracy trend bar chart (last 100 samples, batched in groups)\n'
    '- Latest predictions table with green/red correct/incorrect highlighting\n'
    '- Reset and clear controls\n\n'
    'Processor Scaling Fix: When loading raw CSV data (slow path), the dashboard now '
    'automatically checks for processor.pkl and scales the data before computing metrics. '
    'This prevents the data scale mismatch bug where models trained on scaled data would '
    'receive unscaled inputs, producing incorrect metrics.\n\n'
    'Run: streamlit run src/dashboard/streamlit_app.py'
)

# ============================================================
doc.add_heading('13. API (src/api/main.py) + Streamer (src/api/streamer.py)', level=1)
doc.add_paragraph(
    'FastAPI with REST + WebSocket endpoints (all secured except root and health):\n'
    '- GET / (info) — public\n'
    '- GET /live/status — public (health check)\n'
    '- GET /models — requires auth\n'
    '- POST /predict/{name} — requires auth\n'
    '- GET /live/next/{name} — requires auth, polls next test sample + prediction\n'
    '- WS /ws/live/{name} — requires auth (first message must include username/password)\n'
    '- POST /attack/fgsm, POST /attack/pgd, POST /defense/squeeze, POST /upload.\n\n'
    'Two authentication methods:\n'
    '1. HTTP Basic Auth (username/password from configs/config.py; default: admin / ciciomt2024)\n'
    '2. X-API-Key header (for programmatic clients)\n\n'
    'WebSocket supports control commands: stop, resume, attack, noattack, reset.\n\n'
    'IoMTStreamer class (streamer.py): loads processed test set, yields samples one at a '
    'time with configurable delay, optional FGSM perturbation, and per-sample prediction. '
    'Used by both the REST polling and WebSocket streaming endpoints.\n\n'
    'Real-time performance: XGBoost averages ~2-5ms per prediction, well within the '
    '<100ms requirement for network intrusion detection.\n\n'
    'Run: uvicorn src.api.main:app --reload'
)

# ============================================================
doc.add_heading('14. Deployment', level=1)
doc.add_paragraph(
    'Dockerfile: python:3.10-slim, installs requirements, exposes ports 8000 (API) and 8501 (dashboard).\n\n'
    'docker-compose.yml: Two services (api, dashboard) with shared volumes for data/models/reports.\n\n'
    'Run: docker-compose -f deployment/docker-compose.yml up'
)

# ============================================================
doc.add_heading('15. Tests', level=1)
doc.add_paragraph(
    'test_attacks.py — FGSM and PGD output shape.\n'
    'test_defenses.py — Feature squeezing shape and valid accuracy.\n'
    'test_evaluation.py — Metrics in [0,1] and ASR correctness.\n'
    'test_preprocessing.py — Loader and processor functions.\n\n'
    'Run: python -m pytest tests/ -v'
)

# ============================================================
doc.add_heading('16. How to Run Everything', level=1)
doc.add_paragraph('1. pip install -r requirements.txt')
doc.add_paragraph('2. Place dataset in data/raw/')
doc.add_paragraph('3. python scripts/run_pipeline.py')
doc.add_paragraph('4. streamlit run src/dashboard/streamlit_app.py  (login: admin / ciciomt2024)')
doc.add_paragraph('5. (Optional) uvicorn src.api.main:app --reload')
doc.add_paragraph('6. (Optional) Open src/dashboard/streamlit_app.py → click "Next Sample" for live inference')
doc.add_paragraph('7. (Optional) Connect to ws://localhost:8000/ws/live/xgboost for WebSocket streaming')
doc.add_paragraph('8. (Optional) docker-compose -f deployment/docker-compose.yml up')

# ============================================================
doc.add_heading('17. Results Summary', level=1)
doc.add_heading('Clean Performance', level=2)
t = doc.add_table(rows=4, cols=5)
t.style = 'Light Grid Accent 1'
for j, h in enumerate(['Model', 'Accuracy', 'Precision', 'Recall', 'F1 Score']):
    t.rows[0].cells[j].text = h
data = [
    ['Random Forest', '0.7230', '0.7625', '0.7230', '0.6928'],
    ['XGBoost',       '0.9370', '0.9411', '0.9370', '0.9363'],
    ['LightGBM',      '0.5727', '0.6605', '0.5727', '0.5635'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        t.rows[i+1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph('XGBoost is the best model on clean data with 93.70% accuracy and 93.63% F1.')
doc.add_paragraph(
    'Why XGBoost outperforms:\n'
    '- Boosting vs Bagging: XGBoost sequentially corrects errors of previous trees, capturing '
    'finer decision boundaries for the 51-class problem. Random Forest averages independent trees '
    'and cannot match this granularity with only 30 estimators.\n'
    '- Built-in Regularization: XGBoost applies L1/L2 regularization, preventing overfit while '
    'maintaining capacity. Random Forest has no equivalent and relies on hard depth/split limits.\n'
    '- Level-wise Growth: XGBoost grows trees level by level, which is robust to suboptimal '
    'hyperparameters. LightGBM uses leaf-wise growth that needs deeper trees or more iterations '
    'to converge — with 30 trees at depth 4 it underfits the 51-class space.'
)

doc.add_heading('Attack Impact', level=2)
doc.add_paragraph(
    'At epsilon=0.01, FGSM drops accuracy from 93.70% to 18.38%.\n'
    'PGD drops it to 12.42%.\n\n'
    'Both attacks are devastating. Defenses are essential.'
)

doc.add_heading('Defense Comparison (FGSM)', level=2)
t2 = doc.add_table(rows=5, cols=7)
t2.style = 'Light Grid Accent 1'
for j, h in enumerate(['Defense', 'eps=0', 'eps=0.01', 'eps=0.05', 'eps=0.1', 'eps=0.2', 'eps=0.5']):
    t2.rows[0].cells[j].text = h
ddata = [
    ['None (XGBoost)',       '0.9370', '0.1838', '0.1645', '0.1632', '0.1463', '0.1304'],
    ['Adversarial Training', '0.6134', '0.2603', '0.3394', '0.3803', '0.2798', '0.2373'],
    ['Ensemble',             '0.8495', '0.0827', '0.0727', '0.0585', '0.0508', '0.0363'],
    ['Feature Selection',    '0.1634', '0.1516', '0.1515', '0.0947', '0.0998', '0.0775'],
]
for i, row in enumerate(ddata):
    for j, val in enumerate(row):
        t2.rows[i+1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'Adversarial training provides the most balanced robustness. '
    'The ensemble has high clean accuracy but collapses under attack. '
    'Feature selection (20 features) is too aggressive for 51 classes.'
)

# ============================================================
doc.save('report_thesis/Project_Documentation.docx')
print("Done")
